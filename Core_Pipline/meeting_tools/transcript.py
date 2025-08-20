import argparse
import os
import sys
from datetime import datetime
from typing import Optional, Tuple, List

import requests
from msal import PublicClientApplication, SerializableTokenCache
from dotenv import load_dotenv


GRAPH_BASE_URL = "https://graph.microsoft.com/v1.0"


def _parse_scopes(raw: Optional[str]) -> List[str]:
    if not raw:
        return [
            "User.Read",
            "OnlineMeetings.Read",
            "OnlineMeetingTranscript.Read.All",
        ]
    # Support space- or comma-separated
    separators = [",", " "]
    scopes: List[str] = [raw]
    for sep in separators:
        if sep in raw:
            scopes = [s.strip() for s in raw.replace(",", " ").split() if s.strip()]
            break
    return scopes


def load_config() -> Tuple[str, str, str, List[str], str]:
    load_dotenv()
    tenant_id = os.getenv("TENANT_ID")
    client_id = os.getenv("CLIENT_ID")
    delegated_scopes = _parse_scopes(os.getenv("DELEGATED_SCOPES"))
    cache_path = os.getenv("TOKEN_CACHE_PATH") or os.path.join(
        os.path.expanduser("~"), ".msal", "transcript_tool_cache.bin"
    )

    if not tenant_id or not client_id:
        print("TENANT_ID and CLIENT_ID must be set in .env", file=sys.stderr)
        sys.exit(1)

    authority = f"https://login.microsoftonline.com/{tenant_id}"
    return tenant_id, client_id, authority, delegated_scopes, cache_path


def _ensure_dir(path: str) -> None:
    directory = os.path.dirname(path)
    if directory and not os.path.exists(directory):
        os.makedirs(directory, exist_ok=True)


def acquire_token(authority: str, client_id: str, scopes: List[str], cache_path: str) -> str:
    cache = SerializableTokenCache()
    if os.path.exists(cache_path):
        try:
            cache.deserialize(open(cache_path, "r").read())
        except Exception:
            # If cache is corrupted, start fresh
            cache = SerializableTokenCache()

    app = PublicClientApplication(client_id=client_id, authority=authority, token_cache=cache)

    # Try silent from cache first
    accounts = app.get_accounts()
    if accounts:
        result = app.acquire_token_silent(scopes=scopes, account=accounts[0])
        if result and "access_token" in result:
            return result["access_token"]

    # Device code interactive flow (one-time; will persist to cache)
    flow = app.initiate_device_flow(scopes=scopes)
    if "user_code" not in flow:
        print("Failed to create device code flow.", file=sys.stderr)
        sys.exit(1)
    print(flow["message"])  # instruct user
    result = app.acquire_token_by_device_flow(flow)

    if "access_token" not in result:
        print(f"Authentication failed: {result}", file=sys.stderr)
        sys.exit(1)

    # Persist cache after successful interactive auth
    if cache.has_state_changed:
        try:
            _ensure_dir(cache_path)
            with open(cache_path, "w") as f:
                f.write(cache.serialize())
        except PermissionError:
            # Fallback to a local, writable path within the project directory
            local_cache_path = os.path.join(os.getcwd(), ".msal_cache.bin")
            try:
                with open(local_cache_path, "w") as f:
                    f.write(cache.serialize())
                print(f"Note: TOKEN_CACHE_PATH not writable. Using local cache: {local_cache_path}")
            except Exception as e:
                print(f"Warning: Failed to save token cache ({e}). Continuing without persistent cache.")

    return result["access_token"]


def graph_get(url: str, token: str, params: Optional[dict] = None, headers: Optional[dict] = None) -> requests.Response:
    hdrs = {"Authorization": f"Bearer {token}"}
    if headers:
        hdrs.update(headers)
    response = requests.get(url, headers=hdrs, params=params)
    return response


def resolve_meeting_id(meeting_input: str, token: str) -> Optional[str]:
    # If it's a URL, resolve by joinWebUrl
    if meeting_input.lower().startswith("http://") or meeting_input.lower().startswith("https://"):
        url = f"{GRAPH_BASE_URL}/me/onlineMeetings"
        params = {"$filter": f"joinWebUrl eq '{meeting_input}'"}
        resp = graph_get(url, token, params=params)
        if resp.status_code == 200:
            data = resp.json()
            values = data.get("value", [])
            if values:
                return values[0]["id"]
        else:
            print(f"Failed to search meeting by joinWebUrl: {resp.status_code} {resp.text}", file=sys.stderr)
            return None

    # Try treating it as the Graph onlineMeeting id
    url = f"{GRAPH_BASE_URL}/me/onlineMeetings/{meeting_input}"
    resp = graph_get(url, token)
    if resp.status_code == 200:
        return resp.json().get("id")

    # Try resolving by joinMeetingIdSettings/joinMeetingId (often shown in Teams invites)
    url = f"{GRAPH_BASE_URL}/me/onlineMeetings"
    params = {"$filter": f"joinMeetingIdSettings/joinMeetingId eq '{meeting_input}'"}
    resp = graph_get(url, token, params=params)
    if resp.status_code == 200:
        data = resp.json()
        values = data.get("value", [])
        if values:
            return values[0]["id"]

    print("Could not resolve meeting. Provide a valid meeting Graph id, join URL, or meeting ID.", file=sys.stderr)
    return None


def pick_latest_transcript(transcripts: list) -> Optional[dict]:
    if not transcripts:
        return None

    def parse_dt(item):
        dt = item.get("createdDateTime") or item.get("lastModifiedDateTime")
        try:
            return datetime.fromisoformat(dt.replace("Z", "+00:00")) if dt else datetime.min
        except Exception:
            return datetime.min

    transcripts_sorted = sorted(transcripts, key=parse_dt, reverse=True)
    return transcripts_sorted[0]


def get_meeting_subject(meeting_id: str, token: str) -> str:
    url = f"{GRAPH_BASE_URL}/me/onlineMeetings/{meeting_id}"
    resp = graph_get(url, token)
    if resp.status_code == 200:
        try:
            subject = (resp.json() or {}).get("subject")
            if isinstance(subject, str) and subject.strip():
                return subject.strip()
        except Exception:
            pass
    return meeting_id


def _safe_filename(name: str) -> str:
    safe = "".join(ch if ch.isalnum() or ch in (" ", "-", "_") else "_" for ch in name).strip()
    safe = "_".join(safe.split())
    return safe or "meeting"


def get_transcript_content(meeting_id: str, token: str, fmt: str) -> Tuple[bytes, str]:
    # List transcripts for the meeting
    list_url = f"{GRAPH_BASE_URL}/me/onlineMeetings/{meeting_id}/transcripts"
    list_resp = graph_get(list_url, token)
    if list_resp.status_code != 200:
        print(f"Failed to list transcripts: {list_resp.status_code} {list_resp.text}", file=sys.stderr)
        sys.exit(1)

    items = list_resp.json().get("value", [])
    if not items:
        print("No transcripts available for this meeting.", file=sys.stderr)
        sys.exit(2)

    chosen = pick_latest_transcript(items)
    transcript_id = chosen["id"]

    # Download content
    content_url = f"{GRAPH_BASE_URL}/me/onlineMeetings/{meeting_id}/transcripts/{transcript_id}/content"
    # Graph currently supports text-based transcript content (e.g., text/vtt). If DOCX is requested,
    # fetch VTT and convert locally.
    accept = "text/vtt"
    content_resp = graph_get(content_url, token, headers={"Accept": accept})
    if content_resp.status_code != 200:
        print(f"Failed to download transcript content: {content_resp.status_code} {content_resp.text}", file=sys.stderr)
        sys.exit(1)

    # Always return VTT bytes here; caller handles conversion when needed
    return content_resp.content, "vtt"


def run_tool(meeting_value: str, fmt: str = "vtt") -> str:
    _, client_id, authority, scopes, cache_path = load_config()
    token = acquire_token(authority=authority, client_id=client_id, scopes=scopes, cache_path=cache_path)

    meeting_id = resolve_meeting_id(meeting_value, token)
    if not meeting_id:
        sys.exit(1)

    vtt_bytes, _ = get_transcript_content(meeting_id, token, fmt)

    subject = get_meeting_subject(meeting_id, token)
    base_name = f"{_safe_filename(subject)}_transctipt"
    out_dir = os.path.join(os.getcwd(), "saved_transcripts")
    os.makedirs(out_dir, exist_ok=True)

    if fmt == "vtt":
        out_name = os.path.join(out_dir, f"{base_name}.vtt")
        with open(out_name, "wb") as f:
            f.write(vtt_bytes)
        print(f"Saved transcript to {out_name}")
        return out_name

    # DOCX conversion path
    try:
        from docx import Document  # python-docx
    except Exception:
        out_name = os.path.join(out_dir, f"{base_name}.vtt")
        with open(out_name, "wb") as f:
            f.write(vtt_bytes)
        print("python-docx is not installed. Saved VTT instead. Install with: pip install python-docx")
        print(f"Saved transcript to {out_name}")
        return out_name

    text = vtt_bytes.decode("utf-8", errors="replace")
    doc = Document()
    doc.add_heading("Teams Meeting Transcript", level=1)
    doc.add_paragraph(text)
    out_name = os.path.join(out_dir, f"{base_name}.docx")
    doc.save(out_name)
    print(f"Saved transcript to {out_name}")
    return out_name


def _interactive_main() -> None:
    meeting_value = input("Enter a meeting Graph id, join URL, or joinMeetingId value: ").strip()
    if not meeting_value:
        print("Error: input is required.")
        return
    fmt = input("Choose format (vtt/docx) [vtt]: ").strip().lower() or "vtt"
    if fmt not in ("vtt", "docx"):
        print("Invalid format. Using vtt.")
        fmt = "vtt"
    run_tool(meeting_value, fmt)


def main() -> None:
    # If invoked via menu (no argv passed), fall back to interactive prompts
    if len(sys.argv) == 1:
        _interactive_main()
        return

    parser = argparse.ArgumentParser(description="Download a Microsoft Teams meeting transcript via Microsoft Graph (delegated auth with device code and persistent cache).")
    parser.add_argument("meeting", help="Meeting Graph id, join URL, or joinMeetingId value")
    parser.add_argument("--format", dest="fmt", choices=["vtt", "docx"], default="vtt", help="Output format (default: vtt)")
    args = parser.parse_args()

    run_tool(args.meeting, args.fmt)


if __name__ == "__main__":
    main()
